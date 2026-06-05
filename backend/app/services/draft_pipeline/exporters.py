"""
Document exporters for compiled drafts.

The CTDP pipeline ends with ``ctx.final_draft`` (a markdown string
assembled by the Compiler phase). This module turns that markdown
into three portable formats:

* PDF    — WeasyPrint (HTML+CSS → PDF)
* DOCX   — python-docx (markdown → WordprocessingML)
* LaTeX  — a deterministic article-class skeleton

Design notes
------------
- All three exporters accept the same ``(final_draft, project_id)``
  signature so the router can dispatch uniformly. The ``project_id``
  is currently used only for filename hints in the rendered LaTeX
  metadata and is otherwise ignored — the markdown is the source of
  truth.
- Markdown is parsed once into a lightweight intermediate
  representation (headings + paragraphs + tables) so each exporter
  can render it without re-parsing. Keeping the IR minimal (no
  spans/links) means we don't drag in a full CommonMark parser; a
  block-level pass is enough for the academic paper shape the
  pipeline produces.
- WeasyPrint requires GTK3 / Pango native libraries. On platforms
  where the import fails (notably stock Windows) the import here
  succeeds because the package is installed, but the first
  ``HTML.write_pdf`` call raises. The router catches that and returns
  503 with a clear message so users get an actionable error rather
  than a 500.
"""

from __future__ import annotations

import html as _html
import logging
import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import Iterable, List, Optional, Sequence, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Lightweight markdown → block IR
# ---------------------------------------------------------------------------


@dataclass
class Block:
    """One block-level element of a parsed document."""

    kind: str  # "heading" | "paragraph" | "blockquote" | "table" | "rule"
    text: str = ""
    level: int = 0
    rows: List[List[str]] = field(default_factory=list)


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
_FENCE_RE = re.compile(r"^```")


def _split_table_row(line: str) -> List[str]:
    """Split a markdown table row on unescaped pipes, trimming cells."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip().replace("\\|", "|") for c in line.split("|")]


def parse_markdown_blocks(md: str) -> List[Block]:
    """Parse markdown into a flat list of :class:`Block`.

    The parser is intentionally minimal:

    * ``# … ######`` → ``Block("heading", level=N)``
    * blank-line separated paragraphs → ``Block("paragraph")``
    * ``> …`` → ``Block("blockquote")``
    * ``---``/``***`` → ``Block("rule")``
    * pipe tables → ``Block("table", rows=[[cells…]])``

    Code fences are passed through as paragraph text (their inner
    contents are not re-interpreted). The exporter is expected to
    ``<pre>``/escape them.
    """
    if not md:
        return []
    blocks: List[Block] = []
    lines = md.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block — collect verbatim into a single paragraph
        # marked with the fence info so exporters can render <pre>.
        if _FENCE_RE.match(stripped):
            info = stripped[3:].strip()
            i += 1
            body_lines: List[str] = []
            while i < n and not _FENCE_RE.match(lines[i].strip()):
                body_lines.append(lines[i])
                i += 1
            # consume closing fence if present
            if i < n:
                i += 1
            blocks.append(
                Block(
                    kind="paragraph",
                    text=f"```{(info + ' ') if info else ''}" + "\n".join(body_lines) + "\n```",
                )
            )
            continue

        # Heading
        m = _HEADING_RE.match(line)
        if m:
            blocks.append(Block(kind="heading", level=len(m.group(1)), text=m.group(2).strip()))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^\s*([-*_])\s*\1\s*\1[\s\1]*$", line):
            blocks.append(Block(kind="rule"))
            i += 1
            continue

        # Table — peek at the next line for the separator row.
        if "|" in line and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]):
            rows: List[List[str]] = [_split_table_row(line)]
            i += 2  # skip header + separator
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(_split_table_row(lines[i]))
                i += 1
            blocks.append(Block(kind="table", rows=rows))
            continue

        # Blockquote — collect contiguous ``> `` lines.
        if stripped.startswith(">"):
            q_lines: List[str] = []
            while i < n and lines[i].strip().startswith(">"):
                q_lines.append(lines[i].strip()[1:].lstrip())
                i += 1
            blocks.append(Block(kind="blockquote", text="\n".join(q_lines).strip()))
            continue

        # Blank line — paragraph break.
        if not stripped:
            i += 1
            continue

        # Paragraph — collect until next blank line / structural line.
        para_lines: List[str] = [line.rstrip()]
        i += 1
        while i < n:
            nxt = lines[i]
            nxt_s = nxt.strip()
            if not nxt_s:
                break
            if _HEADING_RE.match(nxt):
                break
            if nxt_s.startswith(">"):
                break
            if "|" in nxt and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]):
                break
            if _FENCE_RE.match(nxt_s):
                break
            if re.match(r"^\s*([-*_])\s*\1\s*\1[\s\1]*$", nxt):
                break
            para_lines.append(nxt.rstrip())
            i += 1
        blocks.append(Block(kind="paragraph", text="\n".join(para_lines).strip()))

    return blocks


# ---------------------------------------------------------------------------
# Inline markdown → plain text + citation handling
# ---------------------------------------------------------------------------


# CTDP uses `[@paper_id]` for in-text citations. We keep them verbatim
# in every output format because the source document is the contract;
# stripping them would lose information the user explicitly added.
_CITATION_RE = re.compile(r"\[@[A-Za-z0-9_:\-.]+\]")


def strip_inline_formatting(s: str) -> str:
    """Reduce inline markdown (bold/italic/code/links) to plain text.

    Citations ``[@id]`` are preserved. Used when the target format
    doesn't support rich text (e.g. plain headings inside LaTeX section
    titles when they accidentally contain a backtick).
    """
    # Strip backtick code spans
    s = re.sub(r"`([^`]+)`", r"\1", s)
    # Strip bold/italic markers
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"__(.+?)__", r"\1", s)
    s = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", s)
    s = re.sub(r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)", r"\1", s)
    # Convert [text](url) → "text (url)"
    s = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", s)
    return s


def render_inline_html(s: str) -> str:
    """Render inline markdown to a small HTML subset safe for both
    WeasyPrint and python-docx's HTML passthrough."""
    return _render_inline_html_impl(s)


def _render_inline_html_impl(s: str) -> str:
    """HTML inline renderer. Kept separate so :func:`render_inline_html`
    stays the public, well-named entry point."""
    out = _html.escape(s, quote=False)
    # Inline code: `text`
    out = re.sub(r"`([^`]+)`", r"<code>\1</code>", out)
    # Bold: **text**
    out = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", out)
    out = re.sub(r"__(.+?)__", r"<strong>\1</strong>", out)
    # Italic: *text* / _text_
    out = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<em>\1</em>", out)
    out = re.sub(r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)", r"<em>\1</em>", out)
    # Links: [text](url)
    out = re.sub(
        r"\[([^\]]+)\]\(([^)]+)\)",
        r'<a href="\2">\1</a>',
        out,
    )
    return out


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------


def to_docx(final_draft: str, project_id: str) -> bytes:
    """Render ``final_draft`` to a DOCX byte string.

    Block-level structure (headings, paragraphs, blockquotes, tables)
    is preserved; inline markdown is reduced to plain text. Citations
    in the ``[@id]`` form are kept verbatim.
    """
    # Import lazily so the rest of the module is usable even when
    # python-docx is missing in some edge deployment.
    from docx import Document
    from docx.shared import Pt

    blocks = parse_markdown_blocks(final_draft)
    doc = Document()

    # Match the academic-paper look-and-feel the rest of the project
    # uses for Word exports: serif body, 11pt, 1.15 line spacing.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    for block in blocks:
        if block.kind == "heading":
            level = max(1, min(block.level, 4))
            text = strip_inline_formatting(block.text)
            doc.add_heading(text, level=level)
        elif block.kind == "paragraph":
            text = strip_inline_formatting(block.text)
            # Detect code-fence paragraphs and render as a quote so
            # the monospace is at least visually distinct.
            if text.startswith("```") and text.endswith("```"):
                inner = text[3:-3]
                if "\n" in inner:
                    first, _, rest = inner.partition("\n")
                    lang = first.strip()
                    body = rest.rstrip()
                else:
                    lang = ""
                    body = inner
                p = doc.add_paragraph()
                run = p.add_run(f"[code{((' ' + lang) if lang else '')}]\n{body}")
                run.italic = True
            else:
                doc.add_paragraph(text)
        elif block.kind == "blockquote":
            text = strip_inline_formatting(block.text)
            doc.add_paragraph(text, style="Intense Quote")
        elif block.kind == "rule":
            doc.add_paragraph("―" * 20)
        elif block.kind == "table" and block.rows:
            cols = max(len(r) for r in block.rows)
            table = doc.add_table(rows=len(block.rows), cols=cols)
            table.style = "Light Grid Accent 1"
            for r_idx, row in enumerate(block.rows):
                for c_idx in range(cols):
                    cell_text = (
                        strip_inline_formatting(row[c_idx])
                        if c_idx < len(row)
                        else ""
                    )
                    table.cell(r_idx, c_idx).text = cell_text

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# LaTeX export
# ---------------------------------------------------------------------------


_LATEX_SPECIAL_CHARS = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def _latex_escape(s: str) -> str:
    """Escape LaTeX special characters. ``\\`` is escaped first so the
    other replacements' backslashes don't get re-escaped."""
    out_chars: List[str] = []
    for ch in s:
        if ch in _LATEX_SPECIAL_CHARS:
            out_chars.append(_LATEX_SPECIAL_CHARS[ch])
        else:
            out_chars.append(ch)
    return "".join(out_chars)


def _latex_inline(s: str) -> str:
    """Inline renderer for LaTeX. Strips inline markdown and escapes
    LaTeX specials, then re-injects citations using ``\\cite``-style
    tokens the article-class ``\\bibliography`` block can pick up."""
    plain = strip_inline_formatting(s)
    escaped = _latex_escape(plain)
    # Map CTDP `[@paper_id]` to a stable cite key derived from the id.
    def _repl(m: re.Match) -> str:
        pid = m.group(0)[2:-1]  # strip [@ and ]
        key = re.sub(r"[^A-Za-z0-9]", "_", pid)
        if not key or not key[0].isalpha():
            key = "ref" + key
        return f"\\cite{{{key}}}"
    return _CITATION_RE.sub(_repl, escaped)


def _latex_render_blocks(blocks: Sequence[Block]) -> str:
    """Render the block IR to a LaTeX body fragment (no preamble)."""
    out: List[str] = []
    for block in blocks:
        if block.kind == "heading":
            level = max(1, min(block.level, 4))
            text = _latex_inline(block.text)
            if level == 1:
                out.append(rf"\section*{{{text}}}")
            elif level == 2:
                out.append(rf"\section*{{{text}}}")
            elif level == 3:
                out.append(rf"\subsection*{{{text}}}")
            else:
                out.append(rf"\subsubsection*{{{text}}}")
        elif block.kind == "paragraph":
            text = block.text
            if text.startswith("```") and text.endswith("```"):
                inner = text[3:-3]
                if "\n" in inner:
                    first, _, rest = inner.partition("\n")
                    lang = first.strip()
                else:
                    lang = ""
                    rest = inner
                out.append(r"\begin{verbatim}")
                out.append(rest.rstrip())
                out.append(r"\end{verbatim}")
            else:
                out.append(_latex_inline(text))
        elif block.kind == "blockquote":
            out.append(r"\begin{quote}")
            out.append(_latex_inline(block.text))
            out.append(r"\end{quote}")
        elif block.kind == "rule":
            out.append(r"\noindent\rule{\linewidth}{0.4pt}")
        elif block.kind == "table" and block.rows:
            cols = max(len(r) for r in block.rows)
            col_spec = "l" * cols
            out.append(rf"\begin{{tabular}}{{{col_spec}}}")
            for r_idx, row in enumerate(block.rows):
                cells = [_latex_inline(c) for c in row]
                # Pad short rows
                cells += [""] * (cols - len(cells))
                out.append(" & ".join(cells) + r" \\")
                if r_idx == 0:
                    out.append(r"\hline")
            out.append(r"\end{tabular}")
        out.append("")
    return "\n".join(out).rstrip() + "\n"


def to_latex(final_draft: str, project_id: str) -> bytes:
    """Render ``final_draft`` to a LaTeX article-class document.

    Returns the document body *without* a ``\\documentclass`` wrapper
    would be the cleanest contract, but every reviewer we asked
    expects to be able to ``pdflatex`` the output without further
    edits, so we ship the full preamble.

    The ``project_id`` is used to seed the ``\\title`` and PDF
    metadata. The article is intentionally library-light (no
    ``\\usepackage{hyperref}``) so it compiles on a vanilla TeX Live
    install. The citation keys are deterministic hashes of the
    original ``[@paper_id]`` so the user can paste a ``.bib`` file
    alongside without us having to mint real BibTeX entries.
    """
    blocks = parse_markdown_blocks(final_draft)

    # Try to lift the first H1 as the document title.
    title = "Compiled Draft"
    body_blocks = list(blocks)
    if body_blocks and body_blocks[0].kind == "heading" and body_blocks[0].level == 1:
        title = strip_inline_formatting(body_blocks[0].text)
        body_blocks = body_blocks[1:]

    safe_id = re.sub(r"[^A-Za-z0-9_-]", "-", project_id or "draft")
    title_escaped = _latex_escape(title)
    body = _latex_render_blocks(body_blocks)

    preamble = (
        r"\documentclass[11pt,a4paper]{article}" + "\n"
        r"\usepackage[utf8]{inputenc}" + "\n"
        r"\usepackage[T1]{fontenc}" + "\n"
        r"\usepackage{lmodern}" + "\n"
        r"\usepackage[margin=1in]{geometry}" + "\n"
        r"\usepackage{parskip}" + "\n"
        r"\title{" + title_escaped + "}\n"
        r"\author{CiteThreads Draft Pipeline}\n"
        r"\date{\today}\n"
        "\n"
        r"\begin{document}" + "\n"
        r"\maketitle" + "\n"
        "\n"
    )
    closing = (
        r"\bibliographystyle{plain}" + "\n"
        r"% Bibliography entries are intentionally omitted: the source "
        "draft uses\n"
        r"% `[@paper_id]` placeholders. Add a matching \begin{thebibliography}"
        "\n"
        r"% block or supply a \bibliography{refs} file at compile time." + "\n"
        r"\nocite{*}" + "\n"
        r"\end{document}" + "\n"
    )
    doc = preamble + body + "\n" + closing
    return doc.encode("utf-8")


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------


_PDF_CSS = """
@page {
  size: A4;
  margin: 1in;
}
body {
  font-family: "Helvetica", "Arial", sans-serif;
  font-size: 11pt;
  line-height: 1.5;
  color: #111;
}
h1, h2, h3, h4 {
  font-family: "Helvetica", "Arial", sans-serif;
  color: #1a1a1a;
  margin-top: 1.4em;
  margin-bottom: 0.5em;
  line-height: 1.25;
}
h1 { font-size: 22pt; border-bottom: 1px solid #ddd; padding-bottom: 0.2em; }
h2 { font-size: 16pt; }
h3 { font-size: 13pt; }
h4 { font-size: 12pt; font-style: italic; }
p { margin: 0 0 0.8em 0; }
blockquote {
  margin: 0.5em 1.2em;
  padding: 0.4em 0.8em;
  border-left: 3px solid #aaa;
  color: #333;
  background: #f7f7f7;
}
code {
  font-family: "Courier New", "Menlo", monospace;
  font-size: 10pt;
  background: #f4f4f4;
  padding: 0 0.2em;
  border-radius: 2px;
}
pre {
  font-family: "Courier New", "Menlo", monospace;
  font-size: 10pt;
  background: #f4f4f4;
  padding: 0.6em 0.8em;
  border-radius: 3px;
  white-space: pre-wrap;
}
table {
  border-collapse: collapse;
  margin: 1em 0;
  width: 100%;
}
th, td {
  border: 1px solid #ccc;
  padding: 0.4em 0.6em;
  text-align: left;
  vertical-align: top;
}
th { background: #f0f0f0; }
hr { border: 0; border-top: 1px solid #ddd; margin: 1.2em 0; }
a { color: #1a5fb4; text-decoration: none; }
"""


def _render_blocks_html(blocks: Sequence[Block]) -> str:
    """Render the block IR to an HTML fragment (no <html>/<body>)."""
    out: List[str] = []
    for block in blocks:
        if block.kind == "heading":
            level = max(1, min(block.level, 4))
            out.append(f"<h{level}>{_render_inline_html_impl(block.text)}</h{level}>")
        elif block.kind == "paragraph":
            text = block.text
            if text.startswith("```") and text.endswith("```"):
                inner = text[3:-3]
                if "\n" in inner:
                    first, _, rest = inner.partition("\n")
                    lang = first.strip()
                else:
                    lang = ""
                    rest = inner
                # Render the code block as <pre><code>.
                lang_attr = f' class="language-{_html.escape(lang)}"' if lang else ""
                code = _html.escape(rest.rstrip())
                out.append(f"<pre><code{lang_attr}>{code}</code></pre>")
            else:
                out.append(f"<p>{_render_inline_html_impl(text)}</p>")
        elif block.kind == "blockquote":
            out.append(f"<blockquote><p>{_render_inline_html_impl(block.text)}</p></blockquote>")
        elif block.kind == "rule":
            out.append("<hr/>")
        elif block.kind == "table" and block.rows:
            cols = max(len(r) for r in block.rows)
            head, *body_rows = block.rows
            head_cells = "".join(
                f"<th>{_render_inline_html_impl(c)}</th>" for c in head
            )
            out.append(
                "<table><thead><tr>" + head_cells + "</tr></thead><tbody>"
            )
            for row in body_rows:
                cells = "".join(
                    f"<td>{_render_inline_html_impl(c) if i < len(row) else ''}</td>"
                    for i, c in enumerate(row)
                )
                # Pad short rows.
                cells += "<td></td>" * (cols - len(row))
                out.append(f"<tr>{cells}</tr>")
            out.append("</tbody></table>")
    return "\n".join(out)


def to_pdf(final_draft: str, project_id: str) -> bytes:
    """Render ``final_draft`` to a PDF byte string via WeasyPrint.

    WeasyPrint requires the GTK3 / Pango native libraries. On Windows
    these are not shipped with the wheel, so the import succeeds but
    the first ``HTML.write_pdf`` call raises ``OSError`` when it
    tries to ``dlopen`` libgobject. The router catches that error and
    returns 503; tests that exercise the import without a working
    native stack should ``pytest.skip`` rather than fail.
    """
    # Imported lazily so the rest of this module is usable even on
    # machines that haven't installed the GTK3 native libraries yet.
    from weasyprint import HTML

    blocks = parse_markdown_blocks(final_draft)
    body_html = _render_blocks_html(blocks)
    safe_id = _html.escape(project_id or "draft")
    full = (
        "<!doctype html>\n"
        '<html lang="en">\n'
        "<head>\n"
        '<meta charset="utf-8">\n'
        f"<title>{safe_id}</title>\n"
        f"<style>{_PDF_CSS}</style>\n"
        "</head>\n"
        f"<body>\n{body_html}\n</body>\n"
        "</html>\n"
    )
    return HTML(string=full, base_url=".").write_pdf()


__all__ = [
    "Block",
    "parse_markdown_blocks",
    "to_docx",
    "to_latex",
    "to_pdf",
    "render_inline_html",
    "strip_inline_formatting",
]
